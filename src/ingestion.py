import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from langchain_community.document_loaders import TextLoader
from langchain.schema import Document
from typing import Dict, Iterable, List, Optional, Tuple, Union

import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document as DocxFile
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from src.manifest import MANIFEST_PATH, load_manifest, now_iso, save_manifest, sha256_file, sha256_text

logger = logging.getLogger(__name__)

# среднее число символов на страницу, ниже которого PDF считается сканом без текстового слоя
OCR_MIN_CHARS_PER_PAGE = 20

# минимальная доля площади текстового блока внутри таблицы,
# при которой блок считается частью таблицы и исключается из обычного текста
TABLE_BLOCK_OVERLAP = 0.5

# типы файлов, обрабатываемые ingestion
RAW_FILE_PATTERNS = ("*.pdf", "*.docx", "*.txt")


def _table_rows_to_markdown(rows: List[List[Optional[str]]]) -> str:
    """Сериализует таблицу (список строк с ячейками) в markdown.

    Объединённые ячейки DOCX/PDF дублируют содержимое — повторы соседних
    ячеек в строке заменяются на пустые. Возвращает "" для вырожденных
    таблиц (меньше двух строк или одного столбца) — такие артефакты
    детектора лучше оставить обычным текстом.
    """

    def fmt(cell: Optional[str]) -> str:
        if not cell:
            return ""
        return re.sub(r"\s+", " ", str(cell)).replace("|", "\\|").strip()

    clean_rows = []
    for row in rows:
        cells = [fmt(c) for c in row]
        if not any(cells):
            continue
        # повтор соседней ячейки слева — след горизонтального объединения
        deduped = [cells[0]]
        for i in range(1, len(cells)):
            deduped.append("" if cells[i] and cells[i] == cells[i - 1] else cells[i])
        clean_rows.append(deduped)

    if len(clean_rows) < 2:
        return ""
    n_cols = max(len(r) for r in clean_rows)
    if n_cols < 2:
        return ""

    lines = []
    for row in clean_rows:
        row = row + [""] * (n_cols - len(row))
        lines.append("| " + " | ".join(row) + " |")
    lines.insert(1, "| " + " | ".join(["---"] * n_cols) + " |")
    return "\n".join(lines)


def ocr_pdf(path: str, dpi: int = 300, lang: str = "rus") -> List[Document]:
    """Извлекает текст из отсканированного PDF через рендер страниц в изображение + Tesseract OCR."""

    pdf = fitz.open(path)
    documents = []
    for page_number in range(pdf.page_count):
        pix = pdf[page_number].get_pixmap(dpi=dpi)
        mode = "RGB" if pix.n < 4 else "RGBA"
        image = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
        text = pytesseract.image_to_string(image, lang=lang)
        documents.append(
            Document(page_content=text, metadata={"source": path, "page": page_number})
        )
    return documents


def _maybe_ocr_pdf(source: str, pages: List[Document], ocr_lang: str = "rus", ocr_dpi: int = 300) -> List[Document]:
    """Заменяет страницы PDF без текстового слоя (скан) на результат OCR."""

    if not pages:
        return pages
    avg_chars = sum(len(d.page_content.strip()) for d in pages) / len(pages)
    if avg_chars < OCR_MIN_CHARS_PER_PAGE:
        logger.info("OCR: %s (текстовый слой отсутствует, %d стр.)", source, len(pages))
        return ocr_pdf(source, dpi=ocr_dpi, lang=ocr_lang)
    return pages


def _pdf_page_content(page: "fitz.Page") -> str:
    """Извлекает текст страницы PDF, сериализуя таблицы в markdown.

    Таблицы находятся детектором PyMuPDF; текстовые блоки, лежащие внутри
    рамки таблицы, исключаются из обычного текста (иначе содержимое ячеек
    дублировалось бы построчной выгрузкой). Блоки и таблицы объединяются
    в порядке следования на странице (по вертикальной координате).
    """

    table_items: List[Tuple[float, str]] = []
    table_rects: List[fitz.Rect] = []
    for table in page.find_tables().tables:
        markdown = _table_rows_to_markdown(table.extract())
        if markdown:
            rect = fitz.Rect(table.bbox)
            table_items.append((rect.y0, markdown))
            table_rects.append(rect)

    if not table_items:
        return page.get_text()

    text_items: List[Tuple[float, str]] = []
    for x0, y0, x1, y1, text, _, block_type in page.get_text("blocks"):
        if block_type != 0 or not text.strip():
            continue
        rect = fitz.Rect(x0, y0, x1, y1)
        in_table = any(
            (rect & table_rect).get_area() / max(rect.get_area(), 1.0)
            > TABLE_BLOCK_OVERLAP
            for table_rect in table_rects
        )
        if not in_table:
            text_items.append((y0, text.strip()))

    items = sorted(text_items + table_items, key=lambda item: item[0])
    return "\n\n".join(content for _, content in items)


def load_pdf(path: str) -> List[Document]:
    """Загружает PDF постранично (PyMuPDF) с таблицами в виде markdown."""

    pdf = fitz.open(path)
    return [
        Document(
            page_content=_pdf_page_content(pdf[page_number]),
            metadata={"source": path, "page": page_number},
        )
        for page_number in range(pdf.page_count)
    ]


def _iter_docx_blocks(docx: "DocxFile") -> Iterable[Union[DocxParagraph, DocxTable]]:
    """Итерирует параграфы и таблицы DOCX в порядке следования в документе."""

    for child in docx.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield DocxParagraph(child, docx)
        elif isinstance(child, CT_Tbl):
            yield DocxTable(child, docx)


def load_docx(path: str) -> List[Document]:
    """Загружает DOCX (python-docx): текст параграфов + таблицы в markdown.

    В отличие от docx2txt сохраняет структуру таблиц — ячейки не
    рассыпаются в плоский текст.
    """

    docx = DocxFile(path)
    blocks = []
    for block in _iter_docx_blocks(docx):
        if isinstance(block, DocxParagraph):
            if block.text.strip():
                blocks.append(block.text)
        else:
            markdown = _table_rows_to_markdown(
                [[cell.text for cell in row.cells] for row in block.rows]
            )
            if markdown:
                blocks.append(markdown)
    return [Document(page_content="\n\n".join(blocks), metadata={"source": path})]


def discover_raw_files(data_dir: str = "data/raw") -> List[Path]:
    """Список исходных файлов (PDF/DOCX/TXT) в каталоге, отсортированный для детерминированности."""

    base = Path(data_dir)
    files: List[Path] = []
    for pattern in RAW_FILE_PATTERNS:
        files.extend(base.rglob(pattern))
    return sorted(set(files))


def load_single_document(path: Path) -> List[Document]:
    """Загружает один исходный файл (постранично для PDF) без очистки/OCR-решения."""

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return load_pdf(str(path))
    if suffix == ".docx" or suffix == ".doc":
        return load_docx(str(path))
    if suffix == ".txt":
        return TextLoader(str(path), encoding="utf-8").load()
    raise ValueError(f"неподдерживаемый тип файла: {suffix}")


def clean_text(text: str) -> str:
    """Очистка текста для русскоязычных нормативных документов.

    Помимо общих артефактов PDF-экстракции убирает колонтитулы КонсультантПлюс
    (шапка "Документ предоставлен...", подвал "надежная правовая поддержка",
    номера страниц, ссылки на www.consultant.ru) — типичные для документов,
    выгруженных с сайта КонсультантПлюс в PDF/DOCX.
    """

    # колонтитулы КонсультантПлюс
    text = re.sub(r'Документ предоставлен КонсультантПлюс\s*\n?', '', text)
    text = re.sub(r'Дата сохранения:\s*\d{2}\.\d{2}\.\d{4}\s*\n?', '', text)
    text = re.sub(r'www\.consultant\.ru\s*\n?', '', text)
    text = re.sub(r'КонсультантПлюс\s*\n?\s*надежная правовая поддержка\s*\n?', '', text)
    text = re.sub(r'Напечатано с сайта[^\n]*\n?', '', text)
    text = re.sub(r'Страница\s+\d+\s+из\s+\d+\s*\n?', '', text)
    # сюда добавить оставшиемся символы, подлежащие удалению

    # убираем дублирующиеся пробелы и переносы
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)

    # КонсультантПлюс дублирует титульный блок документа в начале файла — убираем повторы абзацев
    paragraphs = text.split('\n\n')
    deduped = []
    for p in paragraphs:
        if deduped and p.strip() and deduped[-1].strip() == p.strip():
            continue
        deduped.append(p)
    text = '\n\n'.join(deduped)

    # убираем артефакты нумерации страниц (одиночные числа на отдельной строке)
    text = re.sub(r'\n\s*-?\s*\d+\s*-?\s*\n', '\n', text)
    text = re.sub(r'^\s*\d+\s*\n', '', text)
    text = re.sub(r'\n\s*\d+\s*$', '', text)

    # нормализуем кавычки
    text = text.replace('«', '"').replace('»', '"')

    # убираем мягкий перенос (символ \xad в PDF)
    text = text.replace('\xad', '')

    return text.strip()


def preprocess_documents(documents: List[Document]) -> List[Document]:
    """Предобработка всех документов."""
    for doc in documents:
        doc.page_content = clean_text(doc.page_content)
    # фильтруем пустые документы (в т.ч. отсканированные PDF без текстового слоя)
    return [doc for doc in documents if len(doc.page_content) > 100]


def merge_documents_by_source(documents: List[Document]) -> List[Document]:
    """Объединяет постраничные документы (PDF грузится по страницам) в один документ на файл-источник."""

    grouped: Dict[str, List[Document]] = {}
    for doc in documents:
        source = doc.metadata.get("source", "")
        grouped.setdefault(source, []).append(doc)

    merged = []
    for source, docs in grouped.items():
        docs.sort(key=lambda d: d.metadata.get("page", 0))
        content = "\n\n".join(d.page_content for d in docs)
        metadata = {**docs[0].metadata}
        metadata.pop("page", None)
        metadata["page_count"] = len(docs)
        merged.append(Document(page_content=content, metadata=metadata))

    return merged


def process_single_raw_file(path: Path) -> Document:
    """Полный цикл обработки одного исходного файла: загрузка (+ OCR для сканов PDF) -> очистка -> слияние в один документ.

    Бросает исключение при ошибке загрузки/парсинга или если после очистки
    не осталось текста — вызывающая сторона (run_preprocessing_pipeline)
    перехватывает её на уровне файла, чтобы один битый документ не обрывал
    обработку остальных.
    """

    pages = load_single_document(path)
    if path.suffix.lower() == ".pdf":
        pages = _maybe_ocr_pdf(str(path), pages)
    pages = preprocess_documents(pages)
    if not pages:
        raise ValueError("после очистки не осталось текста (пустой документ или OCR не дал результата)")
    return merge_documents_by_source(pages)[0]


def save_processed_documents(
    documents: List[Document], output_dir: str = "data/processed"
) -> List[Path]:
    """Сохраняет предобработанные документы в виде .txt файлов (один файл на источник)."""

    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    saved_paths = []
    for doc in documents:
        source = doc.metadata.get("source", "unknown")
        stem = Path(source).stem
        out_path = out_dir / f"{stem}.txt"
        out_path.write_text(doc.page_content, encoding="utf-8")
        saved_paths.append(out_path)

    return saved_paths


@dataclass
class IngestResult:
    """Итог стадии ingest: какие исходники обработаны/пропущены/удалены/провалились."""

    processed: List[str] = field(default_factory=list)
    skipped: List[str] = field(default_factory=list)
    failed: List[Tuple[str, str]] = field(default_factory=list)
    deleted: List[str] = field(default_factory=list)


def run_preprocessing_pipeline(
    data_dir: str = "data/raw",
    output_dir: str = "data/processed",
    manifest_path: Path = MANIFEST_PATH,
    force: bool = False,
) -> IngestResult:
    """Фаза 1, инкрементально: обрабатывает только новые/изменившиеся исходники (по манифесту),
    помечает удалённые исходники (файл больше не существует в data/raw) для последующей
    зачистки в Qdrant на стадии индексации, и не даёт ошибке на одном файле оборвать весь прогон.
    """

    manifest = load_manifest(manifest_path)
    raw_files = discover_raw_files(data_dir)
    raw_sources = {str(p) for p in raw_files}
    result = IngestResult()

    # удалённые исходники: были в манифесте, файла больше нет на диске
    for source, entry in manifest.items():
        if source in raw_sources or entry.get("raw_deleted"):
            continue
        processed_path = entry.get("processed_path")
        if processed_path and Path(processed_path).exists():
            Path(processed_path).unlink()
        entry["raw_deleted"] = True
        result.deleted.append(source)
        logger.info("Удалён исходник: %s", source)

    for path in raw_files:
        source = str(path)
        raw_hash = sha256_file(path)
        entry = manifest.get(source, {})
        processed_path = entry.get("processed_path")
        unchanged = (
            not force
            and entry.get("raw_sha256") == raw_hash
            and processed_path
            and Path(processed_path).exists()
        )
        if unchanged:
            result.skipped.append(source)
            continue

        try:
            doc = process_single_raw_file(path)
            out_path = save_processed_documents([doc], output_dir)[0]
            manifest[source] = {
                "raw_sha256": raw_hash,
                "processed_path": str(out_path),
                "processed_sha256": sha256_text(doc.page_content),
                "ingested_at": now_iso(),
                # indexed_sha256/point_ids сохраняются от предыдущего запуска (если были) —
                # index-стадия сама увидит расхождение processed_sha256 и переиндексирует
                "indexed_sha256": entry.get("indexed_sha256"),
                "point_ids": entry.get("point_ids", []),
            }
            result.processed.append(source)
            logger.info("Обработан: %s -> %s", source, out_path)
        except Exception as exc:
            logger.error("Ошибка обработки %s: %s", source, exc)
            result.failed.append((source, str(exc)))

    save_manifest(manifest, manifest_path)
    logger.info(
        "Ingest: обработано %d, без изменений %d, ошибок %d, удалено %d",
        len(result.processed), len(result.skipped), len(result.failed), len(result.deleted),
    )
    return result


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(name)s: %(message)s")
    run_preprocessing_pipeline()
